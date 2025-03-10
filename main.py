import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext
from tkinterdnd2 import DND_FILES, TkinterDnD
import PyPDF2
from docx import Document
from textblob import TextBlob
import re
import os
import threading
from tkinter import messagebox

class SpellCheckerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Korektor Tekstu")
        self.root.geometry("600x400")
        
        # Styl
        style = ttk.Style()
        style.configure("Custom.TButton", padding=10)
        
        # Główny kontener
        main_frame = ttk.Frame(root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Obszar przeciągania
        self.drop_label = ttk.Label(
            main_frame,
            text="Przeciągnij plik PDF lub DOCX tutaj\nlub kliknij przycisk 'Wybierz plik'",
            padding=20,
            relief="solid"
        )
        self.drop_label.pack(fill=tk.X, pady=10)
        
        # Przycisk wyboru pliku
        self.select_button = ttk.Button(
            main_frame,
            text="Wybierz plik",
            command=self.select_file,
            style="Custom.TButton"
        )
        self.select_button.pack(pady=10)
        
        # Pasek postępu
        self.progress = ttk.Progressbar(
            main_frame,
            mode='determinate',
            length=300
        )
        
        # Pole wyników
        self.results_area = scrolledtext.ScrolledText(
            main_frame,
            wrap=tk.WORD,
            height=10
        )
        self.results_area.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Konfiguracja przeciągania i upuszczania
        self.drop_label.drop_target_register(DND_FILES)
        self.drop_label.dnd_bind('<<Drop>>', self.handle_drop)

    def is_polish_word(self, word):
        polish_chars = set('ąćęłńóśźżĄĆĘŁŃÓŚŹŻ')
        return any(char in polish_chars for char in word) or word.isalpha()

    def extract_text_from_pdf(self, file_path):
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                for i, page in enumerate(pdf_reader.pages):
                    text += page.extract_text() + "\n"
                    self.update_progress((i + 1) / total_pages * 50)
        except Exception as e:
            raise Exception(f"Błąd podczas odczytywania pliku PDF: {str(e)}")
        return text

    def extract_text_from_docx(self, file_path):
        text = ""
        try:
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                self.update_progress((i + 1) / total_paragraphs * 50)
        except Exception as e:
            raise Exception(f"Błąd podczas odczytywania pliku DOCX: {str(e)}")
        return text

    def check_spelling(self, text):
        words = re.findall(r'\b\w+\b', text)
        corrections = []
        total_words = len(words)
        
        for i, word in enumerate(words):
            if not self.is_polish_word(word) or word.isnumeric():
                continue
                
            start_idx = max(0, i - 3)
            end_idx = min(len(words), i + 4)
            context = ' '.join(words[start_idx:end_idx])
            
            try:
                word_blob = TextBlob(word)
                corrected = str(word_blob.correct())
                
                if corrected.lower() != word.lower():
                    corrections.append({
                        'error': word,
                        'correction': corrected,
                        'context': context
                    })
            except:
                continue
                
            self.update_progress(50 + (i + 1) / total_words * 50)
        
        return corrections

    def update_progress(self, value):
        self.progress['value'] = value
        self.root.update_idletasks()

    def process_file(self, file_path):
        try:
            if not file_path.lower().endswith(('.pdf', '.docx')):
                raise Exception("Niewspierany format pliku. Wybierz plik PDF lub DOCX.")
            
            self.progress.pack(pady=10)
            self.progress['value'] = 0
            self.results_area.delete(1.0, tk.END)
            self.select_button['state'] = 'disabled'
            
            if file_path.lower().endswith('.pdf'):
                text = self.extract_text_from_pdf(file_path)
            else:
                text = self.extract_text_from_docx(file_path)
            
            if not text.strip():
                raise Exception("Nie udało się odczytać tekstu z pliku")
            
            corrections = self.check_spelling(text)
            
            if not corrections:
                self.results_area.insert(tk.END, "Nie znaleziono błędów ortograficznych!")
            else:
                result_text = f"Znalezione błędy ({len(corrections)}):\n\n"
                for correction in corrections:
                    result_text += f"Błąd: {correction['error']}\n"
                    result_text += f"Poprawnie: {correction['correction']}\n"
                    result_text += f"Kontekst: {correction['context']}\n"
                    result_text += "-" * 50 + "\n"
                self.results_area.insert(tk.END, result_text)
            
        except Exception as e:
            messagebox.showerror("Błąd", str(e))
        
        finally:
            self.progress.pack_forget()
            self.select_button['state'] = 'normal'

    def select_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Dokumenty", "*.pdf;*.docx")]
        )
        if file_path:
            threading.Thread(target=self.process_file, args=(file_path,), daemon=True).start()

    def handle_drop(self, event):
        file_path = event.data
        if file_path:
            # Usuń cudzysłowy z początku i końca ścieżki, jeśli są obecne
            file_path = file_path.strip('"')
            threading.Thread(target=self.process_file, args=(file_path,), daemon=True).start()

    def handle_drag_enter(self, event):
        self.drop_label.configure(relief="sunken")

    def handle_drag_leave(self, event):
        self.drop_label.configure(relief="solid")

if __name__ == '__main__':
    root = TkinterDnD.Tk()
    app = SpellCheckerApp(root)
    root.mainloop()
